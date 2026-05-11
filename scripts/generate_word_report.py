#!/usr/bin/env python3
"""
generate_word_report.py
Generate an executive Word (.docx) report from Allure results + TRX file.

Usage:
    python generate_word_report.py \
        --allure-results allure-results \
        --trx TestResults/results.trx \
        --output WillScot_ExecutiveReport_Build42.docx \
        --build 42 \
        --env Prod

Requires:  pip install python-docx Pillow
"""
import argparse
import json
import os
import sys
import xml.etree.ElementTree as ET
from datetime import datetime, timezone
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.exit("ERROR: python-docx not installed.  Run: pip install python-docx Pillow")


# ── Colour palette ────────────────────────────────────────────────────────────
GREEN  = RGBColor(0x2E, 0x86, 0x48)
RED    = RGBColor(0xC0, 0x39, 0x2B)
ORANGE = RGBColor(0xE6, 0x7E, 0x22)
GREY   = RGBColor(0x7F, 0x8C, 0x8D)
NAVY   = RGBColor(0x1A, 0x35, 0x5E)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)


def _set_cell_bg(cell, hex_color: str):
    """Apply a solid background colour to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _status_color(status: str) -> RGBColor:
    s = status.lower()
    if s in ("passed", "pass"):   return GREEN
    if s in ("failed", "fail"):   return RED
    if s in ("broken",):          return ORANGE
    return GREY


def _fmt_ms(ms) -> str:
    if ms is None:
        return "—"
    try:
        s = int(ms) // 1000
        return f"{s // 60}m {s % 60}s" if s >= 60 else f"{s}s"
    except Exception:
        return "—"


# ── Read allure-results ───────────────────────────────────────────────────────

def load_allure_results(results_dir: str) -> list[dict]:
    """Return list of test result dicts from *-result.json files."""
    results = []
    p = Path(results_dir)
    if not p.exists():
        return results
    for f in sorted(p.glob("*-result.json")):
        try:
            data = json.loads(f.read_text(encoding="utf-8"))
            results.append(data)
        except Exception:
            pass
    return results


def build_summary(results: list[dict]) -> dict:
    counts = {"passed": 0, "failed": 0, "broken": 0, "skipped": 0, "unknown": 0}
    start_ms: list[int] = []
    stop_ms:  list[int] = []

    for r in results:
        status = r.get("status", "unknown").lower()
        counts[status] = counts.get(status, 0) + 1
        if r.get("start"):
            start_ms.append(r["start"])
        if r.get("stop"):
            stop_ms.append(r["stop"])

    total = sum(counts.values())
    start = min(start_ms) if start_ms else None
    stop  = max(stop_ms)  if stop_ms  else None

    duration_ms = (stop - start) if (start and stop) else None

    def ts(ms):
        if ms is None:
            return "—"
        return datetime.fromtimestamp(ms / 1000, tz=timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")

    return {
        "total":       total,
        "passed":      counts.get("passed",  0),
        "failed":      counts.get("failed",  0),
        "broken":      counts.get("broken",  0),
        "skipped":     counts.get("skipped", 0),
        "start":       ts(start),
        "stop":        ts(stop),
        "duration_ms": duration_ms,
        "duration_fmt": _fmt_ms(duration_ms),
        "pass_rate":   f"{round(counts.get('passed',0) / total * 100)}%" if total else "0%",
    }


# ── Read TRX ─────────────────────────────────────────────────────────────────

def load_trx_summary(trx_path: str) -> dict | None:
    if not trx_path or not Path(trx_path).exists():
        return None
    try:
        tree = ET.parse(trx_path)
        root = tree.getroot()
        ns   = {"t": "http://microsoft.com/schemas/VisualStudio/TeamTest/2010"}
        counters = root.find(".//t:Counters", ns)
        if counters is None:
            return None
        return {
            "total":   counters.get("total",   "?"),
            "passed":  counters.get("passed",  "?"),
            "failed":  counters.get("failed",  "?"),
            "skipped": counters.get("notExecuted", "?"),
        }
    except Exception:
        return None


# ── Word document builder ─────────────────────────────────────────────────────

def build_document(results: list[dict], summary: dict, trx: dict | None,
                   build_no: str, env: str, allure_dir: str) -> Document:
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── Title block ───────────────────────────────────────────────────────────
    title = doc.add_heading("WillScot Homepage Automation", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(f"Executive Test Report  ·  Build #{build_no}  ·  Environment: {env}")
    sub_run.font.size  = Pt(11)
    sub_run.font.color.rgb = GREY
    sub_run.font.italic    = True

    doc.add_paragraph()

    # ── Executive summary table ───────────────────────────────────────────────
    doc.add_heading("Test Execution Summary", level=1)
    tbl = doc.add_table(rows=2, cols=7)
    tbl.style = "Table Grid"

    headers = ["Total", "Passed", "Failed", "Broken", "Skipped", "Duration", "Pass Rate"]
    values  = [
        str(summary["total"]),
        str(summary["passed"]),
        str(summary["failed"]),
        str(summary["broken"]),
        str(summary["skipped"]),
        summary["duration_fmt"],
        summary["pass_rate"],
    ]
    bg_colors = ["1A355E", "2E8648", "C0392B", "E67E22", "7F8C8D", "1A355E", "1A355E"]

    for i, (hdr, val, bg) in enumerate(zip(headers, values, bg_colors)):
        hcell = tbl.rows[0].cells[i]
        vcell = tbl.rows[1].cells[i]
        _set_cell_bg(hcell, bg)
        hcell.paragraphs[0].clear()
        hr = hcell.paragraphs[0].add_run(hdr)
        hr.font.bold       = True
        hr.font.color.rgb  = WHITE
        hr.font.size       = Pt(9)
        hcell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        vcell.paragraphs[0].clear()
        vr = vcell.paragraphs[0].add_run(val)
        vr.font.bold = True
        vr.font.size = Pt(11)
        if hdr == "Passed":  vr.font.color.rgb = GREEN
        if hdr == "Failed":  vr.font.color.rgb = RED
        if hdr == "Broken":  vr.font.color.rgb = ORANGE
        vcell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ── Timing block ──────────────────────────────────────────────────────────
    timing = doc.add_paragraph()
    timing.add_run("Start: ").font.bold = True
    timing.add_run(summary["start"])
    timing.add_run("    Stop: ").font.bold = True
    timing.add_run(summary["stop"])
    timing.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ── Detailed results table ────────────────────────────────────────────────
    doc.add_heading("Scenario Results", level=1)

    col_widths = [Inches(3.2), Inches(0.9), Inches(0.8), Inches(2.0)]
    detail_tbl = doc.add_table(rows=1, cols=4)
    detail_tbl.style = "Table Grid"

    hdr_cells = detail_tbl.rows[0].cells
    for i, (text, w) in enumerate(zip(["Scenario", "Status", "Duration", "Tags"], col_widths)):
        _set_cell_bg(hdr_cells[i], "1A355E")
        hdr_cells[i].paragraphs[0].clear()
        r = hdr_cells[i].paragraphs[0].add_run(text)
        r.font.bold      = True
        r.font.color.rgb = WHITE
        r.font.size      = Pt(9)

    for result in sorted(results, key=lambda x: x.get("name", "")):
        name     = result.get("name", result.get("fullName", "Unknown"))
        status   = result.get("status", "unknown")
        dur_ms   = result.get("stop", 0) - result.get("start", 0) if result.get("start") else None
        tags     = [lbl["value"] for lbl in result.get("labels", []) if lbl.get("name") == "tag"]

        row   = detail_tbl.add_row()
        cells = row.cells

        cells[0].paragraphs[0].clear()
        cells[0].paragraphs[0].add_run(name).font.size = Pt(9)

        cells[1].paragraphs[0].clear()
        sr = cells[1].paragraphs[0].add_run(status.upper())
        sr.font.bold      = True
        sr.font.size      = Pt(9)
        sr.font.color.rgb = _status_color(status)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        cells[2].paragraphs[0].clear()
        cells[2].paragraphs[0].add_run(_fmt_ms(dur_ms)).font.size = Pt(9)
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        cells[3].paragraphs[0].clear()
        cells[3].paragraphs[0].add_run(", ".join(tags) if tags else "—").font.size = Pt(8)

    doc.add_paragraph()

    # ── Screenshots from allure-results ──────────────────────────────────────
    png_files = sorted(Path(allure_dir).glob("*-attachment.png"))[:6]  # cap at 6
    if png_files:
        doc.add_heading("Evidence Screenshots", level=1)
        doc.add_paragraph(
            "Screenshots below are taken from the Allure attachments embedded in allure-results. "
            "Full-size images are available in the archived allure-report."
        ).italic = True
        for png in png_files:
            try:
                doc.add_picture(str(png), width=Inches(5.5))
                doc.add_paragraph()
            except Exception:
                pass

    # ── Pipeline Flow ─────────────────────────────────────────────────────────
    doc.add_heading("Pipeline Architecture", level=1)
    flow = doc.add_paragraph()
    flow.style = "List Bullet"
    steps = [
        ("Checkout",          "Source code is cloned from Git SCM (auto-triggered on every push via SCM polling)."),
        ("Build Docker Image","The .NET 8 + Playwright test project is baked into a Docker image (fully cached on rebuild)."),
        ("Run Tests",         f"{summary['total']} BDD scenarios run inside the Docker container — Playwright + Reqnroll + NUnit with 4 parallel workers. Videos, screenshots and traces are volume-mounted directly to the Jenkins workspace."),
        ("Word Report",       "This document is auto-generated from Allure JSON result files + TRX data."),
        ("Allure Report",     "The Allure Jenkins plugin generates an interactive HTML report published on the build page."),
        ("Archive Evidence",  "All artifacts archived: Allure HTML report, raw results, videos (.webm), Playwright traces, TRX file, and this Word report."),
    ]
    for i, (stage, desc) in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"{stage}: ").bold = True
        p.add_run(desc)

    doc.add_paragraph()

    # ── Artifact Links ────────────────────────────────────────────────────────
    doc.add_heading("Artifact Links", level=1)
    artifacts = [
        f"Allure Report        → <Jenkins URL>/job/<JobName>/{build_no}/allure/",
        f"Archived Allure HTML → <Jenkins URL>/job/<JobName>/{build_no}/artifact/allure-report/",
        f"Raw Allure Results   → <Jenkins URL>/job/<JobName>/{build_no}/artifact/allure-results/",
        f"TRX Results File     → <Jenkins URL>/job/<JobName>/{build_no}/artifact/TestResults/results.trx",
        f"Playwright Traces    → Embedded in Allure results as .zip attachments",
        f"Videos (WebM)        → Embedded in Allure results as .webm attachments",
    ]
    for line in artifacts:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(line).font.size = Pt(9)

    doc.add_paragraph()

    # ── Footer note ───────────────────────────────────────────────────────────
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(
        f"Generated automatically by Jenkins pipeline  ·  "
        f"Framework: Playwright 1.44 + Reqnroll 2.2 + NUnit 4  ·  "
        f"Report date: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    fr.font.size       = Pt(8)
    fr.font.color.rgb  = GREY
    fr.font.italic     = True

    return doc


# ── Entry point ───────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Generate Word executive report from Allure results.")
    parser.add_argument("--allure-results", default="allure-results")
    parser.add_argument("--trx",            default="TestResults/results.trx")
    parser.add_argument("--output",         default="WillScot_ExecutiveReport.docx")
    parser.add_argument("--build",          default="0")
    parser.add_argument("--env",            default="Prod")
    args = parser.parse_args()

    print(f"[Word Report] Reading allure-results from: {args.allure_results}")
    results = load_allure_results(args.allure_results)
    print(f"[Word Report] Found {len(results)} result files")

    summary = build_summary(results)
    trx     = load_trx_summary(args.trx)

    print(f"[Word Report] Summary — Total:{summary['total']}  "
          f"Passed:{summary['passed']}  Failed:{summary['failed']}  "
          f"Duration:{summary['duration_fmt']}")

    doc = build_document(results, summary, trx, args.build, args.env, args.allure_results)
    doc.save(args.output)
    print(f"[Word Report] Saved: {args.output}")


if __name__ == "__main__":
    main()
