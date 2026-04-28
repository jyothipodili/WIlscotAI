"""
Generate WillScot Automation — Docker Setup Guide (Word Document)
Run: python generate_docker_doc.py
Output: d:/WilscotProjectAI/WillScot_Docker_Setup_Guide.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

doc = Document()

# ── Page setup ───────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_font(run, bold=False, italic=False, size=11, color=None):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)   # dark blue
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def heading3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p

def normal(text, bold=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, bold=bold, size=11, color=color)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        r.font.size = Pt(11)
    return p

def numbered(text, level=0):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(3)
    for r in p.runs:
        r.font.size = Pt(11)
    return p

def code_block(lines):
    """Light-grey shaded code box."""
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.3)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        r = p.add_run(line)
        r.font.name = "Courier New"
        r.font.size = Pt(10)
        # grey shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F2F2F2")
        pPr.append(shd)

def screenshot_placeholder(label, description=""):
    """Render a bordered table cell as a screenshot placeholder."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(6)
    # border
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for side in ("top","left","bottom","right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    "12")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "2E74B5")
        tcPr.append(border)
    # content
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(f"[ SCREENSHOT: {label} ]")
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    if description:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(description)
        r2.font.size  = Pt(10)
        r2.font.italic = True
        r2.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    doc.add_paragraph()   # spacing after table

def info_table(rows_data, col_widths=None):
    """Two-column key/value table."""
    tbl = doc.add_table(rows=len(rows_data), cols=2)
    tbl.style = "Table Grid"
    if col_widths is None:
        col_widths = [Inches(2), Inches(4)]
    for i, (key, val) in enumerate(rows_data):
        cells = tbl.rows[i].cells
        cells[0].width = col_widths[0]
        cells[1].width = col_widths[1]
        kp = cells[0].paragraphs[0]
        kr = kp.add_run(key)
        kr.bold = True
        kr.font.size = Pt(10)
        vp = cells[1].paragraphs[0]
        vr = vp.add_run(val)
        vr.font.size = Pt(10)
    doc.add_paragraph()

def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    pb.append(bottom)
    pPr.append(pb)

# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("\n\n")

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_p.add_run("WillScot Automation Framework")
tr.bold = True
tr.font.size = Pt(24)
tr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = subtitle_p.add_run("Docker Setup Guide — Step by Step")
sr.bold = True
sr.font.size = Pt(18)
sr.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

doc.add_paragraph()

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_p.add_run(f"Date: {datetime.date.today().strftime('%B %d, %Y')}     |     Version: 1.0     |     Team: QA Automation")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════
heading1("1. Document Overview")
normal("This guide walks a new team member through building, configuring, and running the "
       "WillScot Playwright automation suite inside Docker Desktop on Windows. "
       "Following these steps produces a reproducible, environment-isolated test runner "
       "without any local .NET SDK installation.")

info_table([
    ("Project",        "WillScot Automation Framework"),
    ("Stack",          ".NET 8 · Playwright 1.44 · Reqnroll 2.2 · NUnit 4.1"),
    ("Base Image",     "mcr.microsoft.com/playwright/dotnet:v1.44.0-jammy  (Ubuntu 22.04)"),
    ("Docker Image",   "willscot-automation:latest"),
    ("Container Name", "willscot-tests"),
    ("Report Output",  ".\\reports\\  (allure-results, TestResults, logs)"),
    ("Author",         "QA Automation Team"),
    ("Date",           datetime.date.today().strftime("%B %d, %Y")),
])

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — PRECONDITIONS
# ═══════════════════════════════════════════════════════════════════════════════
heading1("2. Preconditions")
normal("Ensure ALL of the following are satisfied before proceeding.", bold=True)

heading2("2.1  Hardware & OS")
bullet("Windows 10 Pro / Enterprise (Build 19041 or later) or Windows 11")
bullet("Minimum 8 GB RAM  (16 GB recommended — Playwright Chromium uses ~2 GB shared memory)")
bullet("Minimum 20 GB free disk space (base image ~6 GB + build layers)")
bullet("Virtualization enabled in BIOS/UEFI  (check via Task Manager → Performance → CPU → Virtualization: Enabled)")

heading2("2.2  Software Prerequisites")
bullet("Docker Desktop v4.x or later  —  installed and running (Engine status: green in system tray)")
bullet("WSL 2 backend enabled in Docker Desktop Settings → General → 'Use WSL 2 based engine'")
bullet("Git for Windows  (git --version should return a version number)")
bullet("Access to the project Git repository  (clone or pull to local disk)")
bullet("No conflicting containers named willscot-tests already running  (docker ps will show any)")

heading2("2.3  Network & Secrets")
bullet("Internet access to pull the base image from mcr.microsoft.com  (first build only)")
bullet("Optional: TEAMS_WEBHOOK_URL set in environment if Teams notifications are required")
bullet("Optional: ZEPHYR_JWT_TOKEN and JIRA_API_TOKEN set if test-result sync is needed")

heading2("2.4  Project Files")
normal("The following files must exist in the project root:")
code_block([
    "d:\\WilscotProjectAI\\",
    "├── Dockerfile",
    "├── docker-compose.yml",
    "├── .env",
    "├── docker-entrypoint.sh",
    "├── env\\",
    "│   ├── dev.env",
    "│   ├── qa.env",
    "│   └── uat.env",
    "└── WillscotAutomation\\",
    "    └── WillscotAutomation.csproj",
])

hr()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — INSTALL DOCKER DESKTOP
# ═══════════════════════════════════════════════════════════════════════════════
heading1("3. Step 1 — Install Docker Desktop")

heading2("3.1  Download Docker Desktop")
numbered("Open a browser and navigate to:  https://www.docker.com/products/docker-desktop/")
numbered("Click 'Download for Windows'  →  Save the installer  (Docker Desktop Installer.exe ~560 MB).")
numbered("Double-click the installer and accept defaults.  Enable 'Use WSL 2 instead of Hyper-V' when prompted.")
numbered("Click 'OK'  →  installation completes and a restart is required.")
numbered("After restart, Docker Desktop launches automatically.  Wait for the whale icon in the system tray to stop animating — the engine is ready.")

heading2("3.2  Verify Installation")
normal("Open PowerShell or Command Prompt and run:")
code_block([
    "docker --version",
    "# Expected output:",
    "# Docker version 27.x.x, build xxxxxxx",
    "",
    "docker compose version",
    "# Expected output:",
    "# Docker Compose version v2.x.x",
])

normal("Open Docker Desktop → Settings → General and confirm:")
bullet("'Use WSL 2 based engine'  is checked")
bullet("'Start Docker Desktop when you log in'  is checked  (recommended)")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — PROJECT SETUP
# ═══════════════════════════════════════════════════════════════════════════════
heading1("4. Step 2 — Clone / Prepare the Project")

heading2("4.1  Navigate to the Project Root")
normal("Open PowerShell and change to the project directory:")
code_block([
    "cd d:\\WilscotProjectAI",
    "",
    "# Confirm all required files exist:",
    "dir Dockerfile, docker-compose.yml, .env, docker-entrypoint.sh",
])

heading2("4.2  Configure Environment Variables (optional)")
normal("Open .env in a text editor to review defaults.  Change values only if needed:")
code_block([
    "# .env  (default — QA environment, Chromium, headless, all tests)",
    "TEST_ENV=qa",
    "BROWSER=chromium",
    "HEADLESS=true",
    "FILTER=",
    "",
    "# Secrets — fill these on the host or in CI; never commit real values",
    "TEAMS_WEBHOOK_URL=",
    "ZEPHYR_JWT_TOKEN=",
    "JIRA_API_TOKEN=",
])

normal("Available TEST_ENV values:")
info_table([
    ("dev", "Uses appsettings.json.  Browser visible (headless=false).  For local debugging."),
    ("qa",  "Uses appsettings.QA.json.  Headless.  Teams notifications active."),
    ("uat", "Uses appsettings.Stage.json.  Extended timeouts (45–90 s).  Headless."),
], col_widths=[Inches(1.2), Inches(4.8)])

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — BUILD IMAGE
# ═══════════════════════════════════════════════════════════════════════════════
heading1("5. Step 3 — Build the Docker Image")

heading2("5.1  Run docker build")
normal("From the project root  (d:\\WilscotProjectAI)  run:")
code_block([
    "docker build -t willscot-automation:latest .",
])

normal("What happens during the build  (each RUN = a cached layer):")
info_table([
    ("Layer 1", "Pull base image  mcr.microsoft.com/playwright/dotnet:v1.44.0-jammy  (~1 GB, cached after first pull)"),
    ("Layer 2", "apt-get install jq  (JSON tool for Allure config patching)"),
    ("Layer 3", "Copy .csproj  →  dotnet restore  (NuGet packages, cached until .csproj changes)"),
    ("Layer 4", "Copy full source  →  dotnet build -c Release"),
    ("Layer 5", "mkdir -p /app/allure-results /app/logs /app/TestResults"),
    ("Layer 6", "Copy docker-entrypoint.sh  →  strip CRLF  →  chmod +x"),
], col_widths=[Inches(1.2), Inches(4.8)])

normal("Typical first-build time:  8 – 15 minutes  (mostly NuGet + Playwright browser download).")
normal("Subsequent builds with no code change:  under 60 seconds  (all layers cached).")

heading2("5.2  Verify the Image Exists")
normal("After the build completes, verify the image in Docker Desktop or CLI:")
code_block([
    "docker images willscot-automation",
    "",
    "# Expected output:",
    "# REPOSITORY             TAG       IMAGE ID       CREATED        SIZE",
    "# willscot-automation    latest    ea7209573de1   3 days ago     6.09 GB",
])

heading2("5.3  Docker Desktop — Images View  (Screenshot)")
normal("Open Docker Desktop → Images tab.  You should see the image listed as shown below:")
doc.add_paragraph()
screenshot_placeholder(
    "Docker Desktop → Images — willscot-automation:latest",
    "Shows the willscot-automation image (tag: latest, ID: ea7209573de1, Size: 6.09 GB, Created: 3 days ago)"
)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — RUN TESTS
# ═══════════════════════════════════════════════════════════════════════════════
heading1("6. Step 4 — Run the Test Suite")

heading2("6.1  Run All Tests  (QA environment, default)")
code_block([
    "# From d:\\WilscotProjectAI\\",
    "docker compose up",
])

heading2("6.2  Run Tests with Specific Environment")
code_block([
    "# QA  environment (default)",
    "docker compose up",
    "",
    "# UAT environment",
    "TEST_ENV=uat docker compose up",
    "",
    "# DEV environment  (browser visible — use only on a machine with a display)",
    "TEST_ENV=dev docker compose up",
])

heading2("6.3  Run a Subset of Tests  (NUnit Category Filter)")
code_block([
    "# Run smoke tests only",
    "FILTER=smoke docker compose up",
    "",
    "# Run regression suite on UAT",
    "FILTER=regression TEST_ENV=uat docker compose up",
    "",
    "# Run navigation tests on QA with Firefox",
    "FILTER=navigation BROWSER=firefox docker compose up",
])

heading2("6.4  Monitor Container Output")
normal("Docker prints the test runner banner and live test output to the terminal. "
       "A successful run ends with:")
code_block([
    "======================================================",
    "  Test run complete   Exit code : 0",
    "  TRX results  : /app/TestResults/results.trx",
    "  Allure data  : /app/allure-results",
    "  Logs         : /app/logs",
    "======================================================",
])

heading2("6.5  Docker Desktop — Containers View  (Screenshot)")
normal("While the container is running (or after it has exited), open Docker Desktop → "
       "Containers tab.  You will see the containers as shown below:")
doc.add_paragraph()
screenshot_placeholder(
    "Docker Desktop → Containers — willscotprojectai & willscot-tests",
    "Shows 'willscotprojectai' (parent) and 'willscot-tests' (ID: 135b3cc54caf, Image: willscot-au...) — both created 3 days ago"
)

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — TEST REPORTS
# ═══════════════════════════════════════════════════════════════════════════════
heading1("7. Step 5 — Access Test Reports")

normal("After the container exits, reports are written to the host at:")
code_block([
    "d:\\WilscotProjectAI\\reports\\",
    "├── allure-results\\      ← Allure raw JSON (open with Allure CLI or allure-serve)",
    "├── TestResults\\",
    "│   └── results.trx      ← NUnit TRX file (open in Visual Studio or CI)",
    "└── logs\\               ← Custom test logs",
])

heading2("7.1  View Allure Report  (Recommended)")
normal("If Allure CLI is installed on the host:")
code_block([
    "# Install once (requires Java or Node)",
    "npm install -g allure-commandline",
    "",
    "# Serve the report (opens browser automatically)",
    "allure serve reports\\allure-results",
])

heading2("7.2  View TRX in Visual Studio")
normal("Open results.trx in Visual Studio Test Explorer, or use a CI viewer such as "
       "the GitHub Actions TRX report plugin.")

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — ENVIRONMENT FILE REFERENCE
# ═══════════════════════════════════════════════════════════════════════════════
heading1("8. Environment Configuration Reference")

heading2("8.1  env\\qa.env  (QA — default)")
code_block([
    "TEST_ENV=QA",
    "BROWSER=chromium",
    "HEADLESS=true",
    "TEAMS_WEBHOOK_URL=<your-teams-webhook>",
    "ZEPHYR_JWT_TOKEN=",
    "JIRA_API_TOKEN=",
])

heading2("8.2  env\\uat.env  (UAT / Stage)")
code_block([
    "TEST_ENV=uat",
    "BROWSER=chromium",
    "HEADLESS=true",
    "# Extended timeouts — set in appsettings.Stage.json (45–90 s)",
])

heading2("8.3  env\\dev.env  (Local Debugging)")
code_block([
    "TEST_ENV=dev",
    "BROWSER=chromium",
    "HEADLESS=false   # browser visible on local machine",
    "# Teams / Zephyr integration disabled for dev",
])

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — DOCKER FILE REFERENCE
# ═══════════════════════════════════════════════════════════════════════════════
heading1("9. Key File Reference")

heading2("9.1  Dockerfile  (Image Definition)")
code_block([
    "FROM mcr.microsoft.com/playwright/dotnet:v1.44.0-jammy",
    "",
    "USER root",
    "WORKDIR /app",
    "",
    "# Install jq (Allure config patching at runtime)",
    "RUN apt-get update && apt-get install -y --no-install-recommends jq \\",
    "    && rm -rf /var/lib/apt/lists/*",
    "",
    "# NuGet restore (cached layer)",
    "COPY WillscotAutomation/WillscotAutomation.csproj ./",
    "RUN dotnet restore WillscotAutomation.csproj",
    "",
    "# Full source + Release build",
    "COPY WillscotAutomation/ ./",
    "RUN dotnet build WillscotAutomation.csproj -c Release --no-restore",
    "",
    "# Output directories (bind-mounted at runtime)",
    "RUN mkdir -p /app/allure-results /app/logs /app/TestResults",
    "",
    "# Entrypoint script",
    "COPY docker-entrypoint.sh /docker-entrypoint.sh",
    "RUN sed -i 's/\\r$//' /docker-entrypoint.sh && chmod +x /docker-entrypoint.sh",
    "",
    "ENV TEST_ENV=QA BROWSER=chromium HEADLESS=true FILTER=\"\"",
    "ENTRYPOINT [\"/docker-entrypoint.sh\"]",
])

heading2("9.2  docker-compose.yml  (Orchestration)")
code_block([
    "services:",
    "  willscot-tests:",
    "    image: willscot-automation:latest",
    "    container_name: willscot-tests",
    "    env_file:",
    "      - \"env/${TEST_ENV:-qa}.env\"",
    "    environment:",
    "      - TEST_ENV=${TEST_ENV:-QA}",
    "      - BROWSER=${BROWSER:-chromium}",
    "      - HEADLESS=${HEADLESS:-true}",
    "      - FILTER=${FILTER:-}",
    "    volumes:",
    "      - ./reports/allure-results:/app/allure-results",
    "      - ./reports/TestResults:/app/TestResults",
    "      - ./reports/logs:/app/logs",
    "    shm_size: \"2gb\"          # required for Playwright/Chromium",
    "    security_opt:",
    "      - no-new-privileges:true",
    "    restart: \"no\"",
])

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 10 — COMMON COMMANDS QUICK REFERENCE
# ═══════════════════════════════════════════════════════════════════════════════
heading1("10. Quick Reference — Common Commands")

tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Table Grid"
hdr = tbl.rows[0].cells
hdr[0].paragraphs[0].add_run("Command").bold = True
hdr[1].paragraphs[0].add_run("Purpose").bold = True

commands = [
    ("docker build -t willscot-automation:latest .", "Build the Docker image from source"),
    ("docker compose up", "Run all tests (QA, chromium, headless)"),
    ("TEST_ENV=uat docker compose up", "Run tests against UAT environment"),
    ("FILTER=smoke docker compose up", "Run smoke tests only"),
    ("FILTER=regression TEST_ENV=qa docker compose up", "Run regression suite on QA"),
    ("BROWSER=firefox docker compose up", "Run tests with Firefox browser"),
    ("docker images willscot-automation", "List built images"),
    ("docker ps -a", "List all containers (including stopped)"),
    ("docker logs willscot-tests", "View container log output"),
    ("docker rmi willscot-automation:latest", "Remove the image (forces full rebuild)"),
    ("docker compose down", "Stop and remove the container"),
    ("docker system prune -f", "Clean up unused images/containers (frees disk)"),
]

for cmd, purpose in commands:
    row = tbl.add_row()
    cp = row.cells[0].paragraphs[0]
    cr = cp.add_run(cmd)
    cr.font.name = "Courier New"
    cr.font.size = Pt(9)
    pp = row.cells[1].paragraphs[0]
    pp.add_run(purpose).font.size = Pt(10)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 11 — TROUBLESHOOTING
# ═══════════════════════════════════════════════════════════════════════════════
heading1("11. Troubleshooting")

info_table([
    ("Docker build fails — cannot pull base image",
     "Check internet access.  Corporate proxy?  Set HTTPS_PROXY in Docker Desktop Settings → Resources → Proxies."),
    ("'Error: no such file or directory' for entrypoint",
     "Ensure docker-entrypoint.sh was saved with Unix line endings (LF not CRLF). The Dockerfile strips CRLF automatically via sed."),
    ("Container exits immediately with code 1",
     "Run 'docker logs willscot-tests' to see the test runner output.  One or more tests may have failed — check TestResults\\results.trx."),
    ("No reports generated after run",
     "Ensure the reports\\ directories exist on the host before running: mkdir -p reports\\allure-results reports\\TestResults reports\\logs"),
    ("'Port already in use' error",
     "This test runner has no exposed ports.  If another service conflicts, check docker ps for leftover containers."),
    ("Image size too large",
     "The 6 GB image includes Playwright browser binaries.  This is expected.  Use .dockerignore to exclude local build artifacts from the build context."),
    ("Tests pass locally but fail in Docker",
     "Confirm TEST_ENV, BROWSER, and appsettings config match.  Check env\\qa.env values, especially base URLs and credentials."),
], col_widths=[Inches(2.5), Inches(3.5)])

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 12 — POSTCONDITIONS
# ═══════════════════════════════════════════════════════════════════════════════
heading1("12. Postconditions")
normal("After a successful Docker setup and test run, verify the following:", bold=True)

heading2("12.1  Image Verification")
bullet("docker images willscot-automation  shows tag latest with size ~6 GB")
bullet("Image ID is consistent across machines built from the same source commit")

heading2("12.2  Container Behavior")
bullet("Container willscot-tests runs, executes all (or filtered) tests, and exits with code 0 (pass) or non-zero (failures)")
bullet("Container is removed automatically after each run  (restart: no in compose)")
bullet("No dangling containers remain  (docker ps -a should show willscot-tests with status Exited)")

heading2("12.3  Report Output")
bullet("reports\\allure-results\\  contains Allure JSON files")
bullet("reports\\TestResults\\results.trx  contains NUnit TRX test results")
bullet("reports\\logs\\  contains run-specific log files")
bullet("All report directories are on the host machine and persist after the container exits")

heading2("12.4  Cleanup  (Optional)")
normal("To free disk space after verifying reports:")
code_block([
    "# Remove stopped containers",
    "docker container prune -f",
    "",
    "# Remove the image when you want a clean rebuild",
    "docker rmi willscot-automation:latest",
    "",
    "# Full cleanup (all unused images + containers + networks)",
    "docker system prune -f",
])

heading2("12.5  Team Handoff Checklist")
tbl2 = doc.add_table(rows=1, cols=2)
tbl2.style = "Table Grid"
hdr2 = tbl2.rows[0].cells
hdr2[0].paragraphs[0].add_run("Checklist Item").bold = True
hdr2[1].paragraphs[0].add_run("Status").bold = True
items = [
    "Docker Desktop installed and engine running",
    "WSL 2 backend enabled",
    "Project cloned/pulled to local disk",
    "docker build succeeds without errors",
    "docker compose up runs all tests",
    "reports\\ directory populated with results",
    "Allure report viewed successfully",
    "TEAMS_WEBHOOK_URL set  (if Teams notifications required)",
    "Team member can reproduce build independently",
]
for item in items:
    r = tbl2.add_row()
    r.cells[0].paragraphs[0].add_run(item).font.size = Pt(10)
    r.cells[1].paragraphs[0].add_run("☐  Pass  /  Fail").font.size = Pt(10)

doc.add_paragraph()

# ── Footer note ───────────────────────────────────────────────────────────────
hr()
footnote = doc.add_paragraph()
footnote.alignment = WD_ALIGN_PARAGRAPH.CENTER
fn = footnote.add_run("WillScot Automation Framework  |  QA Team  |  Confidential  |  "
                       + datetime.date.today().strftime("%B %d, %Y"))
fn.font.size  = Pt(9)
fn.font.italic = True
fn.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

# ── Save ──────────────────────────────────────────────────────────────────────
out = r"d:\WilscotProjectAI\WillScot_Docker_Setup_Guide.docx"
doc.save(out)
print(f"Document saved: {out}")
