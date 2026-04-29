from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

# ── Helpers ───────────────────────────────────────────────────────────────────
def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def add_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_step(doc, text):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    # light-gray shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F2F2F2')
    pPr.append(shd)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

def add_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

# ── Title ─────────────────────────────────────────────────────────────────────
add_title(doc, "WillScot Automation – Jenkins CI/CD Setup and Execution Flow")
doc.add_paragraph()

# ── 1. Jenkins Setup ──────────────────────────────────────────────────────────
add_heading(doc, "1) Jenkins Setup")
add_step(doc, "Install Jenkins on Windows VM")
add_step(doc, "Verify Jenkins service:")
add_code(doc, "sc query jenkins")
add_step(doc, "Open Jenkins URL:")
add_code(doc, "http://localhost:8080/")
add_step(doc, "Verify Jenkins is running:")
add_code(doc, "curl -I http://localhost:8080")

# ── 2. Jenkins Login Details ──────────────────────────────────────────────────
add_heading(doc, "2) Jenkins Login Details")
add_bullet(doc, "Jenkins Username:  santhi")
add_bullet(doc, "Jenkins API Token: 7c3e7e0cbe544413b27026c57542c865")
add_bullet(doc, "Jenkins URL:       http://localhost:8080/")

# ── 3. Download Jenkins CLI ───────────────────────────────────────────────────
add_heading(doc, "3) Download Jenkins CLI")
add_step(doc, "Download CLI jar:")
add_code(doc, "curl -L -o jenkins-cli.jar http://localhost:8080/jnlpJars/jenkins-cli.jar")

# ── 4. Install Jenkins Plugins ────────────────────────────────────────────────
add_heading(doc, "4) Install Jenkins Plugins (CLI)")
add_step(doc, "Run:")
add_code(doc, "java -jar jenkins-cli.jar -s http://localhost:8080/ -http -auth santhi:7c3e7e0cbe544413b27026c57542c865 install-plugin allure-jenkins-plugin nunit github-branch-source workflow-aggregator")
add_step(doc, "Restart Jenkins:")
add_code(doc, "java -jar jenkins-cli.jar -s http://localhost:8080/ -http -auth santhi:7c3e7e0cbe544413b27026c57542c865 safe-restart")

# ── 5. GitHub Webhook Setup ───────────────────────────────────────────────────
add_heading(doc, "5) GitHub Webhook Setup")
add_step(doc, "Open GitHub Repo → Settings → Webhooks")
add_step(doc, "Click Add Webhook")
add_step(doc, "Enter Payload URL:")
add_code(doc, "http://10.20.30.40:8080/github-webhook/")
add_step(doc, "Select Content Type: application/json")
add_step(doc, "Select event: Just the push event")
add_step(doc, "Click Add webhook")
add_step(doc, "Verify delivery is successful")

# ── 6. Jenkins Pipeline Job Creation ─────────────────────────────────────────
add_heading(doc, "6) Jenkins Pipeline Job Creation")
add_step(doc, "Create Pipeline job name:  WillScot-Automation")
add_step(doc, "Configure Git repository URL")
add_step(doc, "Select Branch: main")
add_step(doc, "Configure pipeline to use Jenkinsfile from repo")

# ── 7. Jenkinsfile Pipeline Stages ────────────────────────────────────────────
add_heading(doc, "7) Jenkinsfile Pipeline Stages")
for stage in ["Checkout", "Restore", "Build", "Install Playwright Browsers",
              "Run Tests", "Publish NUnit Results", "Publish Allure Report"]:
    add_bullet(doc, stage)

# ── 8. Local Execution ────────────────────────────────────────────────────────
add_heading(doc, "8) Local Execution (End-to-End Flow)")

add_para(doc, "Step 1: Restore and Build")
add_code(doc, "cd d:\\WilscotProjectAI\\WillscotAutomation\ndotnet restore\ndotnet build --configuration Release --no-restore")

add_para(doc, "Step 2: Install Playwright Browser")
add_code(doc, 'cd WillscotAutomation\npowershell -NonInteractive -ExecutionPolicy Bypass -File "bin\\Release\\net8.0\\playwright.ps1" install chromium')

add_para(doc, "Step 3: Run All Tests")
add_code(doc, "dotnet test --no-build --configuration Release --settings WillscotAutomation.runsettings -- NUnit.NumberOfTestWorkers=1")

add_para(doc, "Step 4: Generate and Open Allure Report Locally")
add_code(doc, "allure generate allure-results --clean -o allure-report\nallure open allure-report")

# ── 9. Git Commit and Push Flow ───────────────────────────────────────────────
add_heading(doc, "9) Git Commit and Push Flow")
add_code(doc, 'cd d:\\WilscotProjectAI\ngit status\ngit add .\ngit commit -m "WillScot automation - demo run"\ngit push origin main')

# ── 10. Jenkins Auto Trigger Flow ─────────────────────────────────────────────
add_heading(doc, "10) Jenkins Auto Trigger Flow")
add_step(doc, "Push code to GitHub main branch")
add_step(doc, "Jenkins build starts automatically")
add_step(doc, "Monitor pipeline execution in Jenkins UI")

# ── 11. Jenkins Build Trigger Using CLI ───────────────────────────────────────
add_heading(doc, "11) Jenkins Build Trigger Using CLI")
add_code(doc, 'java -jar /tmp/jenkins-cli.jar -s http://localhost:8080 -auth "santhi:7c3e7e0cbe544413b27026c57542c865" build "WillScot-Automation" -v -f 2>&1')

# ── 12. Run Jenkins CI Test Command ───────────────────────────────────────────
add_heading(doc, "12) Run Jenkins CI Test Command")
add_code(doc, 'dotnet test --no-build --configuration Release --settings WillscotAutomation.runsettings --logger "trx;LogFileName=jenkins-results.trx" --logger "nunit;LogFileName=nunit-results.xml" --results-directory TestResults -- NUnit.NumberOfTestWorkers=1')

# ── 13. Download Jenkins Artifacts ────────────────────────────────────────────
add_heading(doc, "13) Download Jenkins Artifacts")
add_step(doc, "Open Jenkins build number")
add_step(doc, "Go to Build Artifacts")
add_step(doc, "Download:")
add_bullet(doc, "allure-report.zip")
add_bullet(doc, "allure-summary.json")

# ── 14. Open Allure Report from ZIP ───────────────────────────────────────────
add_heading(doc, "14) Open Allure Report from ZIP")
add_step(doc, "Extract allure-report.zip")
add_step(doc, "Run:")
add_code(doc, "python -m http.server 8085")
add_step(doc, "Open:")
add_code(doc, "http://localhost:8085")

# ── Save ──────────────────────────────────────────────────────────────────────
out = r"d:\WilscotProjectAI\WillScot_Jenkins_CICD_Setup.docx"
doc.save(out)
print(f"Saved: {out}")
